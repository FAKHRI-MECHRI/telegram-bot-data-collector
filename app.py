import logging
from telegram import Update
from telegram.ext import ( Application,CommandHandler,MessageHandler,filters,ConversationHandler,ContextTypes,)

import asyncio
from docx import Document
from docx.shared import Inches
import os
from datetime import datetime

# ===== Configuration =====
logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# Conversation states
NAME, PHONE_NUMBER, LOCATION = range(3)

# Document settings
DOCX_FILE = "user_data.docx"
TABLE_HEADERS = ["Timestamp", "Name", "PhoneNumber", "Location"]

# ===== DOCX Handling =====
def init_docx():
    """Initialize the Word document with table structure"""
    if not os.path.exists(DOCX_FILE):
        doc = Document()
        table = doc.add_table(rows=1, cols=len(TABLE_HEADERS))
        table.style = "Table Grid"
        
        # Set headers
        for i, header in enumerate(TABLE_HEADERS):
            table.cell(0, i).text = header
            table.cell(0, i).width = Inches(1.5)
        
        doc.save(DOCX_FILE)

def save_to_docx(user_data):
    """Append user data to the Word document"""
    doc = Document(DOCX_FILE)
    table = doc.tables[0]
    
    # Add new row
    row = table.add_row().cells
    row[0].text = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    row[1].text = user_data.get("name", "")
    row[2].text = user_data.get("phone_number", "")
    row[3].text = user_data.get("location", "")
    
    doc.save(DOCX_FILE)

# ===== Bot Handlers =====
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.message.reply_text("Welcome! Please enter name:")
    return NAME

async def get_name(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data["name"] = update.message.text
    await update.message.reply_text(f"ok {update.message.text}! enter phone number?")
    return PHONE_NUMBER

async def get_phone_number(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data["phone_number"] = update.message.text
    await update.message.reply_text("Great! Set address?")
    return LOCATION

async def get_location(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data["location"] = update.message.text
    
    # Save to Word document
    try:
        save_to_docx(context.user_data)
        await update.message.reply_text(
            "✅ Your information has been saved!\n"
            f"Name: {context.user_data['name']}\n"
            f"PhoneNumber: {context.user_data['phone_number']}\n"
            f"Location: {context.user_data['location']}"
        )
    except Exception as e:
        logger.error(f"Failed to save data: {e}")
        await update.message.reply_text("⚠️ Failed to save your data. Please try again.")
    
    return ConversationHandler.END

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.message.reply_text("Operation cancelled.")
    return ConversationHandler.END

# ===== Main Application =====
def main() -> None:
    # Initialize document
    init_docx()
    
    # Create Application
    application = Application.builder().token("your token").build()
    
    # Setup conversation handler
    conv_handler = ConversationHandler(
        entry_points=[CommandHandler("start", start)],
        states={
            NAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_name)],
            PHONE_NUMBER: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_phone_number)],
            LOCATION: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_location)],
        },
        fallbacks=[CommandHandler("cancel", cancel)],
    )
    
    application.add_handler(conv_handler)
    
    # Start bot
    application.run_polling()

if __name__ == "__main__":
    asyncio.run(main())